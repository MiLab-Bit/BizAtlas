from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document


def export_credit_docx(payload: dict[str, Any], out_path: str | Path) -> Path:
    path = Path(out_path)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(str(payload.get("title") or "企业信用评估报告"), level=0)

    company = payload.get("company") or {}
    doc.add_paragraph(
        f"企业：{company.get('name', '—')}　行业：{company.get('industry') or '—'}　"
        f"等级：{payload.get('grade')}　得分：{payload.get('score')}"
    )
    doc.add_paragraph(str(payload.get("headline") or ""))

    for section in payload.get("sections") or []:
        doc.add_heading(str(section.get("title") or ""), level=1)
        if section.get("body"):
            doc.add_paragraph(str(section["body"]))
        for b in section.get("bullets") or []:
            doc.add_paragraph(str(b), style="List Bullet")

    doc.add_paragraph("— 完 —")
    doc.save(str(path))
    return path


def export_onepager_docx(payload: dict[str, Any], out_path: str | Path) -> Path:
    """Thin wrapper: treat onepager as a short credit-like doc."""
    credit_like = {
        "title": payload.get("title") or "一页风险摘要",
        "company": payload.get("company"),
        "headline": payload.get("headline"),
        "grade": payload.get("grade"),
        "score": payload.get("score"),
        "sections": [
            {
                "title": "结论",
                "body": payload.get("headline"),
                "bullets": [f"等级 {payload.get('grade')}", f"得分 {payload.get('score')}"],
            },
            {
                "title": "五维风险",
                "body": "",
                "bullets": [
                    f"{d.get('id')}: {d.get('score')}" for d in (payload.get("dimensions") or [])
                ],
            },
            {
                "title": "Top 风险点",
                "body": "",
                "bullets": [
                    f"[{h.get('severity')}] {h.get('message')} — {h.get('explain')}"
                    for h in (payload.get("top_risks") or [])
                ],
            },
            {
                "title": "数据说明",
                "body": payload.get("disclaimer"),
                "bullets": [str(payload.get("data_quality") or {})],
            },
        ],
    }
    return export_credit_docx(credit_like, out_path)
